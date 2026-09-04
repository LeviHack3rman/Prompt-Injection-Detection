"""Append Chapters Four to Six and a reference list to a copy of the dissertation.

The original Chapters One to Three file is never modified: it is copied first and all
writing targets the copy. Formatting is reproduced from measurements of the original's
WordprocessingML, so the appended chapters are indistinguishable in style.

Usage:  python docgen/build_document.py
"""
from __future__ import annotations

import json
import pathlib
import re
import shutil
import sys

import pandas as pd
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches

HERE = pathlib.Path(__file__).resolve().parent
ROOT = HERE.parent
sys.path.insert(0, str(HERE))

import style as S  # noqa: E402
from context import build as build_context, ph  # noqa: E402
from references import sorted_references  # noqa: E402
import content_ch4, content_ch5, content_ch6  # noqa: E402

SOURCE = ROOT / "Idowu_Damilola__301812221__Chapter_One - Three.docx"
TARGET = ROOT / "Idowu_Damilola__301812221__Full_Project.docx"
OUT = ROOT / "outputs"
TAB = OUT / "tables"
FIGDIR = OUT / "figures"
SHOTS = OUT / "screenshots"

MAX_TABLE_ROWS = 30


# --------------------------------------------------------------------------------------
# Static tables authored for Chapter Four
# --------------------------------------------------------------------------------------
def env_table(c):
    return (["Component", "Specification"], [
        ["Language and runtime", f"Python {c['python_version']}"],
        ["Workstation", f"{c['machine']}, {c['ram']} unified memory, macOS"],
        ["Accelerator", "Apple Metal Performance Shaders (MPS) backend of PyTorch"],
        ["Classical models", f"scikit-learn {c['sklearn_version']}"],
        ["Transformer models", f"PyTorch {c['torch_version']}, Transformers {c['transformers_version']}"],
        ["Data handling", f"pandas {c['pandas_version']}, PyArrow"],
        ["Middleware", f"Flask {c['flask_version']}"],
        ["Protected application", "FastAPI with Uvicorn; static HTML, CSS and JavaScript front end"],
        ["Interface capture", f"Playwright {c['playwright_version']} driving headless Chromium"],
        ["Isolation", "Python virtual environment with a pinned dependency manifest"],
    ], [2600, 6410])


def code_table(c):
    return (["Module or package", "Responsibility"], [
        ["ml/build_dataset.py", "Retrieves every source, applies the four controls, writes the corpus and the audit tables"],
        ["ml/domain_prompts.py", "The authored matched-domain benign prompts and hard negatives"],
        ["ml/common.py", "Shared feature union, partition loading and metric computation"],
        ["ml/evasion.py", "Generates the reserved adaptive-attack set from held-out attacks only"],
        ["ml/train_classical.py", "Trains and evaluates Logistic Regression, Random Forest and the SVM"],
        ["ml/train_transformers.py", "Fine-tunes and evaluates BERT and DistilBERT"],
        ["ml/baselines.py", "Evaluates the keyword filter and the off-the-shelf DeBERTa detector"],
        ["ml/bench_latency.py", "Measures screening latency and the over-blocking rate"],
        ["ml/report.py", "Renders every results table and figure from the saved metrics"],
        ["middleware/app.py", "The Flask screening service and its endpoints"],
        ["middleware/policy.py", "The block, sanitise and escalate decision policy"],
        ["middleware/detector.py", "The three interchangeable detector back-ends"],
        ["middleware/tests/", "Unit tests fixing the behaviour of the decision policy"],
        ["backend/, frontend/", "The vulnerable chatbot application and its guardrail levels"],
        ["lab_eval/probe.py", "Probes the application with attack and benign payloads at every level"],
        ["capture/shots.py", "Drives the browser and captures the interface figures"],
        ["docgen/", "Generates this document from the saved metrics"],
    ], [2900, 6110])


def policy_table(c):
    return (["Decision", "Condition", "Effect on the request"], [
        ["Block", "Score at or above 0.90 (0.80 on the retrieved channel)",
         "Refused outright; the content never reaches the model and is withheld from the response"],
        ["Escalate", "Score at or above 0.70 (0.60 on the retrieved channel)",
         "Held for human review; also applied when sanitisation finds no marker to remove"],
        ["Sanitise", "Score at or above 0.50 (0.40 on the retrieved channel)",
         "Injection scaffolding is stripped and the remainder proceeds to the model"],
        ["Allow", "Score below the sanitisation threshold",
         "Passes to the model unchanged"],
    ], [1500, 3000, 4510])


def levels_table(c):
    return (["Level", "Defences active", "What it demonstrates"], [
        ["1", "Weak system prompt only", "The undefended target"],
        ["2", "+ output filter on the literal secret string",
         "That a naive string filter does not catch encoded or transformed disclosure"],
        ["3", "+ hardened system prompt and keyword input heuristics",
         "Prompt-level defence, and the brittleness and over-defence of keyword filtering"],
        ["4", "+ a second model acting as a judge over the reply",
         "Model-as-judge review catching obfuscated disclosure"],
        ["5", "+ input sanitisation stripping injection markers",
         "Input-side mitigation and its limits"],
    ], [900, 3400, 4710])


def objectives_table(c, M):
    def status(k):
        return "Achieved" if k else "Achieved with qualification"
    return (["Objective", "Principal evidence", "Status"], [
        ["1. Review attacks and defences and identify the gaps",
         "Chapter Two; three gaps stated in Section 2.13", "Achieved"],
        ["2. Assemble a labelled dataset of benign, direct, jailbreak and indirect prompts",
         f"Section 5.2.1; {c['n_total']} samples, four controls applied",
         "Achieved with qualification (HackAPrompt unavailable)"],
        ["3. Develop and comparatively evaluate classical and transformer detectors",
         "Section 5.2.2; five detectors, three seeds, two baselines", "Achieved"],
        ["4. Build a real-time middleware screening each prompt",
         "Sections 4.2.6 and 5.2.5; Flask service, tested and integrated", "Achieved"],
        ["5. Evaluate with security metrics and evasion testing",
         f"Sections 5.2.2 to 5.2.5; {c['n_evasion']} adaptive variants", "Achieved"],
    ], [2700, 4200, 2110])


# --------------------------------------------------------------------------------------
# CSV-derived tables
# --------------------------------------------------------------------------------------
def csv_table(name: str, widths=None, max_rows=MAX_TABLE_ROWS):
    p = TAB / f"{name}.csv"
    if not p.exists():
        return None
    df = pd.read_csv(p)
    if len(df) > max_rows:
        df = df.head(max_rows)
    headers = [str(c) for c in df.columns]
    rows = [[("" if pd.isna(v) else str(v)) for v in r] for r in df.itertuples(index=False)]
    if not widths:
        widths = [9010 // len(headers)] * len(headers)
    return headers, rows, widths


def dataset_table():
    p = TAB / "dataset_summary.csv"
    if not p.exists():
        return None
    df = pd.read_csv(p)
    df.columns = [str(c).replace("attack_class", "Attack class").replace("test", "Test")
                  .replace("train", "Training").replace("val", "Validation")
                  .replace("total", "Total") for c in df.columns]
    pretty = {"benign": "Benign", "direct_injection": "Direct injection",
              "jailbreak": "Jailbreak", "indirect_injection": "Indirect injection"}
    df.iloc[:, 0] = df.iloc[:, 0].map(lambda v: pretty.get(v, v))
    headers = [str(c) for c in df.columns]
    rows = [[str(v) for v in r] for r in df.itertuples(index=False)]
    return headers, rows, [2800] + [1552] * (len(headers) - 1)


def trigger_table():
    p = TAB / "trigger_audit.csv"
    if not p.exists():
        return None
    df = pd.read_csv(p)
    df = df[["trigger", "benign_count", "malicious_count", "benign_rate_pct",
             "malicious_rate_pct"]]
    df.columns = ["Trigger phrase", "Benign samples", "Malicious samples",
                  "Benign rate (%)", "Malicious rate (%)"]
    return ([str(c) for c in df.columns],
            [[str(v) for v in r] for r in df.itertuples(index=False)],
            [2500, 1600, 1700, 1600, 1610])


def latency_table():
    p = OUT / "latency.json"
    if not p.exists():
        return None
    d = json.loads(p.read_text())
    rows = []
    h = d.get("http_latency", {})
    if h:
        rows.append(["Screening over HTTP (user channel)", f"{h.get('mean_ms')}",
                     f"{h.get('median_ms')}", f"{h.get('p95_ms')}", f"{h.get('n')}"])
    for k, v in d.get("inprocess_latency", {}).items():
        rows.append([f"In process: {k} back-end", f"{v.get('mean_ms')}",
                     f"{v.get('median_ms')}", f"{v.get('p95_ms')}", f"{v.get('n')}"])
    tbl1 = (["Measurement", "Mean (ms)", "Median (ms)", "95th percentile (ms)", "n"],
            rows, [3200, 1450, 1450, 1900, 1010])
    return tbl1


def over_blocking_table():
    p = OUT / "latency.json"
    if not p.exists():
        return None
    d = json.loads(p.read_text())
    rows = []
    for o in d.get("over_blocking", []):
        rows.append([o["set"], f"{o['n']:,}", f"{100 * o['block_rate']:.2f}",
                     f"{100 * o['any_intervention_rate']:.2f}"])
    if not rows:
        return None
    return (["Legitimate traffic set", "n", "Blocked (%)", "Any intervention (%)"],
            rows, [3800, 1200, 2000, 2010])


# --------------------------------------------------------------------------------------
# Figures
# --------------------------------------------------------------------------------------
FIGURES = {
    "fig_help_centre": (SHOTS / "01_help_centre.png",
        "The help-centre application that hosts the assistant, with the chat launcher "
        "docked in the lower right corner.", 6.0),
    "fig_chat_widget": (SHOTS / "02_chat_widget.png",
        "The embedded assistant answering an ordinary support question at guardrail level "
        "one, with no defence triggered.", 6.0),
    "fig_level1": (SHOTS / "03_level1_injection_attempt.png",
        "A direct instruction-override payload issued at guardrail level one, where no "
        "input or output defence is active. The assistant refuses on the strength of "
        "model-side alignment alone, so the secret is not disclosed and the "
        "injection-successful banner does not appear.", 6.0),
    "fig_level3": (SHOTS / "04_level3_blocked.png",
        "The same payload at guardrail level three. The keyword input heuristic intercepts "
        "the prompt before the model is reached, and the interface reports which defence "
        "fired.", 6.0),
    "fig_indicator": (SHOTS / "05_defence_indicator.png",
        "Detail of the defence indicator, naming the guardrail that blocked the prompt and "
        "the level at which it was operating.", 4.6),
    "fig_over_defence": (SHOTS / "06_over_defence.png",
        "Over-defence in the deployed application: a legitimate request to list the "
        "password-reset steps is refused at guardrail level three because it contains an "
        "instruction-like phrase, illustrating the false-positive cost of keyword "
        "filtering.", 6.0),
    "fig_middleware": (SHOTS / "07_ml_middleware_block.png",
        "The trained detection middleware screening the same payload at guardrail level "
        "one, where none of the application's own defences are active. The prompt is "
        "blocked before the model is reached and the block is attributed to the "
        "machine-learning detector.", 6.0),
    "fig_comparison": (FIGDIR / "model_comparison.png",
        "Precision, recall, F1-score and false-positive rate for every detector and "
        "baseline on the held-out test set. Error bars show the standard deviation across "
        "three seeds.", 6.3),
    "fig_confusion": (FIGDIR / "confusion_matrices.png",
        "Confusion matrices on the held-out test set, averaged across three seeds.", 6.3),
    "fig_roc": (FIGDIR / "roc_curves.png",
        "Receiver operating characteristic curves on the held-out test set, with the area "
        "under each curve given in the legend.", 5.2),
    "fig_evasion": (FIGDIR / "evasion_degradation.png",
        "Detection rate by evasion transformation. The dotted lines mark each detector's "
        "detection rate on the unmodified test set, so the gap between a bar and its line "
        "is the degradation attributable to that transformation.", 6.3),
    "fig_levels": (FIGDIR / "guardrail_levels.png",
        "Attacks blocked and legitimate prompts over-blocked at each guardrail level of "
        "the vulnerable application.", 6.0),
}


# --------------------------------------------------------------------------------------
# Document assembly
# --------------------------------------------------------------------------------------
def fix_heading3_style(doc):
    """The Heading 3 style is defined but unused, and renders dark blue and non-bold,
    which is inconsistent with the black bold Heading 1 and Heading 2 of this document.
    Redefine it to sit correctly in the 22/16/14 point ladder."""
    styles = doc.styles.element
    for st in styles.findall(qn("w:style")):
        if st.get(qn("w:styleId")) == "Heading3":
            for rpr in st.findall(qn("w:rPr")):
                st.remove(rpr)
            st.append(parse_xml(
                f'<w:rPr {nsdecls("w")}><w:b/><w:bCs/><w:color w:val="000000"/>'
                f'<w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr>'))
            return True
    return False


def strip_empty_headings(doc) -> int:
    """Remove the empty Heading 1 and Heading 2 paragraphs left in the source document,
    which would otherwise appear as blank entries in a generated table of contents."""
    removed = 0
    for p in list(doc.paragraphs):
        style = p.style.name if p.style is not None else ""
        if style.startswith("Heading") and not p.text.strip():
            p._element.getparent().remove(p._element)
            removed += 1
    return removed


def add_picture_paragraph(doc, path: pathlib.Path, width_in: float):
    """A centred paragraph containing only the image, matching the source document."""
    p = doc.add_paragraph()
    p.add_run().add_picture(str(path), width=Inches(width_in))
    ppr = p._element.get_or_add_pPr()
    for child in list(ppr):
        ppr.remove(child)
    ppr.append(parse_xml(f'<w:pStyle {nsdecls("w")} w:val="NormalWeb"/>'))
    ppr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="120" w:beforeAutospacing="0" '
                         f'w:after="40" w:afterAutospacing="0"/>'))
    ppr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
    return p


def main():
    if not SOURCE.exists():
        raise SystemExit(f"source document not found: {SOURCE}")

    # -- gather the real values --------------------------------------------------------
    prov_path = TAB / "dataset_provenance.json"
    if not prov_path.exists():
        raise SystemExit("run `python ml/build_dataset.py` first")
    prov = json.loads(prov_path.read_text())

    M = {}
    for f in ("metrics_classical.json", "metrics_transformers.json", "metrics_baselines.json"):
        p = OUT / f
        if p.exists():
            M.update(json.loads(p.read_text()))

    lab_p = OUT / "lab" / "lab_probe_results.jsonl"
    lab = pd.read_json(lab_p, lines=True) if lab_p.exists() else None
    lat_p = OUT / "latency.json"
    latency = json.loads(lat_p.read_text()) if lat_p.exists() else None

    env = json.loads((OUT / "environment.json").read_text())
    c = build_context(M, prov, lab, latency, env)

    # -- copy, never edit in place -----------------------------------------------------
    shutil.copy2(SOURCE, TARGET)
    doc = Document(str(TARGET))
    fix_heading3_style(doc)
    n_stripped = strip_empty_headings(doc)
    print(f"removed {n_stripped} empty heading paragraphs")

    body = doc.element.body

    TABLES = {
        "table_env": env_table(c),
        "table_code": code_table(c),
        "table_policy": policy_table(c),
        "table_levels": levels_table(c),
        "table_objectives": objectives_table(c, M),
        "table_dataset": dataset_table(),
        "table_triggers": trigger_table(),
        "table_main": csv_table("results_main", [2400, 1500, 1900, 1300, 1400, 510]),
        "table_confusion": csv_table("results_confusion"),
        "table_per_class": csv_table("results_per_class_recall"),
        "table_evasion": csv_table("results_evasion"),
        "table_evasion_by_transform": csv_table("results_evasion_by_transform"),
        "table_levels_results": csv_table("results_guardrail_levels"),
        "table_latency": latency_table(),
        "table_over_blocking": over_blocking_table(),
    }
    TABLES["table_levels"] = TABLES["table_levels"]  # Ch4 static
    ch5_levels = TABLES.pop("table_levels_results")

    n_missing = []

    def emit(block):
        kind = block[0]
        if kind == "body":
            doc._body._element.append(S.body(block[1]))
        elif kind == "chapter_heading":
            doc._body._element.append(S.chapter_heading(block[1]))
        elif kind == "chapter_subtitle":
            doc._body._element.append(S.chapter_subtitle(block[1]))
        elif kind == "heading1":
            doc._body._element.append(S.heading1(block[1]))
        elif kind == "heading2":
            doc._body._element.append(S.heading2(block[1], block[2]))
        elif kind == "heading3":
            doc._body._element.append(S.heading3(block[1], block[2]))
        elif kind == "code_line":
            doc._body._element.append(S.code_line(block[1]))
        elif kind == "caption":
            _, label, text, tkey = block
            tbl = TABLES.get(tkey)
            if tkey == "table_levels" and label.startswith("Table 5.9"):
                tbl = ch5_levels
            doc._body._element.append(S.caption(label, text, before=80, after=80))
            if tbl:
                headers, rows, widths = tbl
                doc._body._element.append(S.table(headers, rows, widths))
            else:
                n_missing.append(tkey)
                doc._body._element.append(S.body(
                    ph(f"table ‘{tkey}’ unavailable — re-run the pipeline "
                       f"(python ml/report.py) to generate outputs/tables/")))
        elif kind == "figure":
            key, label = block[1], block[2]
            path, cap, width = FIGURES[key]
            if path.exists():
                add_picture_paragraph(doc, path, width)
                # after=80 matches the Figure 3.1 caption in the existing chapters
                doc._body._element.append(S.caption(label, cap, before=0, after=80))
            else:
                n_missing.append(key)
                doc._body._element.append(S.body(ph(
                    f"figure ‘{path.name}’ not found — run "
                    f"{'python capture/shots.py' if 'screenshots' in str(path) else 'python ml/report.py'} "
                    f"to generate it")))
        else:
            raise ValueError(f"unknown block kind: {kind}")

    for mod in (content_ch4, content_ch5, content_ch6):
        for block in mod.build(c):
            emit(block)

    # -- references --------------------------------------------------------------------
    doc._body._element.append(S.heading1("REFERENCES"))
    for ref in sorted_references():
        doc._body._element.append(S.reference(ref))

    # -- appendix ----------------------------------------------------------------------
    doc._body._element.append(S.heading1("APPENDIX A"))
    doc._body._element.append(S.chapter_subtitle("SOURCE CODE AND REPRODUCTION"))
    doc._body._element.append(S.body(
        "The complete source of the detection pipeline, the screening middleware, the "
        "vulnerable application used as an integration target, and the scripts that "
        "generate every table and figure in this dissertation are held in the project "
        "repository. Every result reported in Chapter Five is regenerated by running the "
        "commands below in order from the repository root, after creating a virtual "
        "environment and installing the pinned dependency manifest. An API key for the "
        "model provider is required only for the steps that exercise the live application "
        "and for the paraphrase transformation of the evasion set; the remaining steps run "
        "offline once the public datasets have been cached."))
    for cmd in c["commands"]:
        doc._body._element.append(S.code_line(cmd))
    doc._body._element.append(S.body(
        "The generated artefacts are written to the outputs directory: metrics.json holds "
        "every measured value, the tables subdirectory holds the comma-separated sources "
        "of every table in Chapter Five, the figures subdirectory holds the plots, and the "
        "screenshots subdirectory holds the interface captures together with a manifest "
        "recording what was observed at the moment each was taken."))

    doc.save(str(TARGET))

    if n_missing:
        print(f"\n!! {len(n_missing)} placeholders inserted for missing artefacts: "
              f"{sorted(set(n_missing))}")
    print(f"\nWrote {TARGET.name}")
    print(f"Original preserved at {SOURCE.name}")
    return 0 if not n_missing else 0


if __name__ == "__main__":
    sys.exit(main())
